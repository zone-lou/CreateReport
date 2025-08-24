from docx import Document
from shutil import copyfile
test_date={
    "宗地代码":"330782102236JC00588",
    "权利人" : "朱逢凯、朱锦青、朱智建",
    "身份证号":"330782199402050810",
    "联系电话":"13705798101",
    "地址":"义亭镇江岸村木桥18幢8号",
    "所有权":"义乌市义亭镇江岸村农民集体",
    "土地权属依据":"义亭镇农建字[2019]第002号",
    "其它权属依据":"义亭镇（街）农处字[2025]第5号",
    "共有情况":"共同共有",
    "建成年份":"2008",
    "批准面积":"120",
    "占地面积":"120.22",
    "总层数":"5",
    "所在层":"1-5",
    "建筑面积(地上)":"382.99",
    "建筑面积(地下)":"130",
    "总建筑面积":"512.99",
    "地上规划面积":"375.88",
    "总规划面积":"495.88",
    "地上超占面积":"7.11",
    "地下超占面积":"10",
    "合计超占":"17.11",
    "权利类型":"宅基地使用权",
    "权利性质":"批准拨用",
    "批准用途":"072农村宅基地",
    }



##读取 word 文档

path=r"/Users/louzeyu/PycharmProjects/CreateReport/template/不动产测量报告.docx"
out_path=r"/Users/louzeyu/PycharmProjects/CreateReport/不动产.docx"
def replace_text(doc_path,out_path,data):

    copyfile(doc_path,out_path)
    doc = Document(out_path)
    for para in doc.paragraphs:
        replace_in_paragraph(para, data)
    for table in doc.tables:  # 如果占位符在表格中，也需要替换
        replace_in_table(table, data)

    doc.save(out_path)


def replace_in_paragraph(paragraph, data):
    """替换单个段落中的占位符（保留格式）"""
    paragraph.text
    for run in paragraph.runs:
        original_text = run.text
        if not original_text:
            continue

        # 遍历所有占位符键
        for key, value in data.items():
            placeholder = "["+key+"]"  # 占位符格式 {{key}}
            if placeholder in original_text:
                # 直接替换文本，保留原格式（run.font 等属性不变）
                run.text = original_text.replace(placeholder, str(value))

def replace_in_table(table, data):
    """替换单个表格中的占位符（保留格式）"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # replace_in_paragraph(paragraph, data)
                for key, value in data.items():
                    if '[' + key + ']' in paragraph.text:
                        paragraph.text = paragraph.text.replace('[' + key + ']', str(value))

replace_text(path,out_path,test_date)